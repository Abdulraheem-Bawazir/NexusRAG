from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument

from app.models.document import Document


def load_docx(file_path: str | Path) -> Document:
    """Load a DOCX file and normalize it into a Document."""

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if not path.is_file():
        raise ValueError(f"Expected a file, got: {path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(
            f"Expected a .docx file, got: {path.suffix or 'no extension'}"
        )

    docx = DocxDocument(path)

    paragraphs = [
        paragraph.text.strip()
        for paragraph in docx.paragraphs
        if paragraph.text.strip()
    ]

    if not paragraphs:
        raise ValueError(
            f"No extractable text found in DOCX: {path.name}"
        )

    text = "\n\n".join(paragraphs)

    return Document(
        id=str(uuid4()),
        text=text,
        source=path.name,
        file_type="docx",
        metadata={
            "source_path": str(path.resolve()),
            "paragraph_count": len(paragraphs),
        },
    )